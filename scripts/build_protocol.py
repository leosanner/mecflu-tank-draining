from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCUMENT_PATH = ROOT / "docs" / "mecflu-experimental-protocol.docx"


def remove_template_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_padding(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_width(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_w = table_pr.first_child_found_in("w:tblW")
    if table_w is None:
        table_w = OxmlElement("w:tblW")
        table_pr.append(table_w)
    total = sum(width.twips for width in widths)
    table_w.set(qn("w:w"), str(int(total)))
    table_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(int(width.twips)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_padding(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_table_text(table, header=True):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if header and row_index == 0:
                        run.bold = True
        if header and row_index == 0:
            set_repeat_table_header(row)


def set_paragraph_text(paragraph, text, parse_subscripts=False):
    if not parse_subscripts:
        paragraph.add_run(text)
        return
    for part in re.split(r"(_[\wΔ]+)", text):
        if not part:
            continue
        run = paragraph.add_run(part[1:] if part.startswith("_") else part)
        run.font.subscript = part.startswith("_")


def add_table(document, headers, rows, widths, subscript_columns=()):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Modern Paper"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for values in rows:
        cells = table.add_row().cells
        for column, (cell, text) in enumerate(zip(cells, values)):
            if column in subscript_columns:
                cell.text = ""
                set_paragraph_text(cell.paragraphs[0], text, parse_subscripts=True)
            else:
                cell.text = text
    set_table_width(table, widths)
    format_table_text(table)
    document.add_paragraph("")
    return table


def add_body(document, text):
    paragraph = document.add_paragraph(style="Block Text")
    set_paragraph_text(paragraph, text, parse_subscripts=True)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_heading(document, text, level=2):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    set_paragraph_text(paragraph, text, parse_subscripts=True)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    set_paragraph_text(paragraph, text, parse_subscripts=True)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def math_element(tag):
    return OxmlElement(f"m:{tag}")


def math_nodes(*parts):
    nodes = []
    for part in parts:
        if isinstance(part, str):
            run = math_element("r")
            text = math_element("t")
            text.text = part
            run.append(text)
            nodes.append(run)
        elif isinstance(part, (list, tuple)):
            nodes.extend(math_nodes(*part))
        else:
            nodes.append(part)
    return nodes


def math_normal(text):
    run = math_element("r")
    properties = math_element("rPr")
    properties.append(math_element("nor"))
    run.append(properties)
    text_element = math_element("t")
    text_element.text = text
    run.append(text_element)
    return run


def math_container(tag, content):
    element = math_element(tag)
    for node in math_nodes(content):
        element.append(node)
    return element


def math_subscript(base, subscript):
    element = math_element("sSub")
    element.append(math_container("e", base))
    element.append(math_container("sub", subscript))
    return element


def math_superscript(base, superscript):
    element = math_element("sSup")
    element.append(math_container("e", base))
    element.append(math_container("sup", superscript))
    return element


def math_fraction(numerator, denominator):
    element = math_element("f")
    element.append(math_container("num", numerator))
    element.append(math_container("den", denominator))
    return element


def math_radical(content):
    element = math_element("rad")
    properties = math_element("radPr")
    hide_degree = math_element("degHide")
    hide_degree.set(qn("m:val"), "1")
    properties.append(hide_degree)
    element.append(properties)
    element.append(math_container("e", content))
    return element


def math_delimiter(content, begin="(", end=")"):
    element = math_element("d")
    properties = math_element("dPr")
    begin_character = math_element("begChr")
    begin_character.set(qn("m:val"), begin)
    end_character = math_element("endChr")
    end_character.set(qn("m:val"), end)
    properties.extend([begin_character, end_character])
    element.append(properties)
    element.append(math_container("e", content))
    return element


def math_bar(content):
    element = math_element("bar")
    properties = math_element("barPr")
    position = math_element("pos")
    position.set(qn("m:val"), "top")
    properties.append(position)
    element.append(properties)
    element.append(math_container("e", content))
    return element


def math_nary(character, lower, upper, content):
    element = math_element("nary")
    properties = math_element("naryPr")
    symbol = math_element("chr")
    symbol.set(qn("m:val"), character)
    limit_location = math_element("limLoc")
    limit_location.set(qn("m:val"), "subSup")
    properties.extend([symbol, limit_location])
    element.append(properties)
    element.append(math_container("sub", lower))
    element.append(math_container("sup", upper))
    element.append(math_container("e", content))
    return element


def math_sub(base, subscript):
    return math_subscript(base, subscript)


def math_symbol(base, subscript=None):
    return math_subscript(base, subscript) if subscript is not None else base


def math_function(name, argument):
    return math_nodes(name, math_delimiter(argument))


def add_equation(document, latex_source, expression):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph._p.set(qn("w:rsidR"), "00000000")
    math_paragraph = math_element("oMathPara")
    math_object = math_element("oMath")
    for node in math_nodes(expression):
        math_object.append(node)
    math_paragraph.append(math_object)
    paragraph._p.append(math_paragraph)
    _ = latex_source  # Retained in the generator as the editable LaTeX source.
    return paragraph


def add_metadata_line(document, label, value):
    paragraph = document.add_paragraph(style="Block Text")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    if label:
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
    paragraph.add_run(value)


def build_protocol():
    document = Document(DOCUMENT_PATH)
    remove_template_body(document)

    properties = document.core_properties
    properties.title = "Protocolo Experimental - Tempo de esvaziamento de um tanque"
    properties.subject = "Etapa 2 do Projeto Integrado de Mecânica dos Fluidos"
    properties.author = "Leonardo Leme Sanner; Eleonora Chofi"
    properties.keywords = "mecânica dos fluidos, Bernoulli, esvaziamento de tanque, protocolo experimental"

    document.add_paragraph("Protocolo Experimental", style="Title")
    document.add_paragraph(
        "Tempo de esvaziamento de um tanque",
        style="Subtitle",
    )
    add_metadata_line(document, "Disciplina", "Mecânica dos Fluidos")
    add_metadata_line(document, "Professor", "Júlio César")
    add_metadata_line(document, "Integrantes", "Leonardo Leme Sanner - RA 169.328")
    add_metadata_line(document, "", "Eleonora Chofi - RA 169.278")
    add_metadata_line(document, "Etapa", "2 - Entrega do Protocolo Experimental")
    add_metadata_line(document, "Data", "1 de junho de 2026")
    document.add_page_break()

    add_heading(document, "1. Objetivo")
    add_body(
        document,
        "Determinar experimentalmente o tempo de descida do nível da água entre duas alturas "
        "definidas em uma garrafa PET e comparar o resultado com o tempo previsto pelo modelo "
        "quase-estacionário obtido pela aplicação da equação de Bernoulli e da continuidade.",
    )
    add_body(
        document,
        "Como objetivo complementar, será filmada uma execução para estimar a evolução temporal "
        "da altura da coluna d'água, h(t), permitindo comparar a curva experimental com o modelo "
        "teórico e discutir os efeitos não contemplados pelo modelo ideal.",
    )

    add_heading(document, "2. Materiais")
    add_table(
        document,
        ["Material", "Finalidade"],
        [
            ["Garrafa PET transparente de 2 L", "Tanque experimental"],
            ["Água", "Fluido de trabalho"],
            ["Régua graduada", "Medição da altura da coluna d'água"],
            ["Cronômetro de celular", "Medição do tempo de esvaziamento"],
            ["Celular com câmera", "Filmagem de uma execução para obtenção de h(t)"],
            ["Paquímetro ou instrumento equivalente", "Medição do diâmetro do orifício"],
            ["Recipiente graduado", "Calibração da relação entre volume e altura"],
            ["Termômetro, se disponível", "Registro da temperatura da água"],
            ["Marcador permanente ou fita adesiva", "Indicação dos níveis inicial, final e intermediários"],
            ["Recipiente auxiliar", "Coleta da água descarregada"],
            ["Suporte ou superfície estável", "Fixação da garrafa na posição vertical"],
        ],
        [Cm(6.2), Cm(10.2)],
    )

    add_heading(document, "3. Montagem experimental")
    add_body(
        document,
        "A garrafa PET será posicionada verticalmente, sem tampa, sobre uma superfície estável. "
        "Será produzido um orifício lateral circular próximo à base, com diâmetro nominal de "
        "aproximadamente 5 mm e bordas regulares. O diâmetro real será medido após a produção do "
        "furo. A água descarregada será direcionada a um recipiente auxiliar.",
    )
    add_body(
        document,
        "Uma régua será fixada externamente à garrafa. A referência h = 0 será definida no centro "
        "do orifício. A altura h corresponde à distância vertical entre essa referência e a superfície "
        "livre. Para reduzir o efeito da geometria irregular da PET, a análise será limitada a uma faixa "
        "de paredes aproximadamente cilíndricas.",
    )
    add_body(document, "Serão identificados na garrafa:")
    add_bullet(document, "o nível inicial h_i;")
    add_bullet(document, "o nível final h_f, mantido acima do regime terminal e, como referência inicial, com h_f/d ≥ 10;")
    add_bullet(document, "níveis intermediários espaçados regularmente, preferencialmente a cada 2 cm ou 3 cm.")

    add_heading(document, "4. Procedimento experimental")
    steps = [
        "Selecionar uma faixa aproximadamente cilíndrica da garrafa PET para definir os níveis inicial e final.",
        "Produzir um furo circular limpo, sem tubo acoplado e sem rebarbas; fotografá-lo com escala e medir dois diâmetros perpendiculares.",
        "Registrar a posição vertical do centro do orifício e fixar a régua externamente à garrafa.",
        "Medir o diâmetro interno médio da faixa analisada ou calibrar sua área efetiva pela relação entre variação de volume e variação de altura.",
        "Tampar provisoriamente o orifício, preencher a garrafa acima do nível inicial, retirar a tampa da garrafa e aguardar a estabilização da água.",
        "Liberar o orifício antes de a superfície livre alcançar a marca superior.",
        "Iniciar a cronometragem quando a superfície livre cruzar o nível inicial e encerrá-la na passagem pelo nível final.",
        "Repetir o procedimento pelo menos três vezes, preferencialmente cinco, mantendo os mesmos níveis inicial e final.",
        "Calcular a média e a dispersão dos tempos obtidos.",
        "Filmar preferencialmente todas as execuções, com a câmera fixa e posicionada perpendicularmente à régua.",
        "Analisar o vídeo em intervalos regulares e registrar os pares (t, h).",
        "Construir os gráficos e comparar os resultados experimentais com o modelo teórico.",
    ]
    for step in steps:
        add_number(document, step)

    add_heading(document, "5. Grandezas medidas e calculadas")
    add_table(
        document,
        ["Símbolo", "Grandeza", "Unidade", "Obtenção"],
        [
            ["d", "Diâmetro do orifício", "m", "Medição direta"],
            ["d_x, d_y", "Diâmetros perpendiculares do orifício", "m", "Medição direta"],
            ["h_i", "Altura inicial da coluna d'água", "m", "Medição direta"],
            ["h_f", "Altura final da coluna d'água", "m", "Medição direta"],
            ["Δt_j", "Duração medida na repetição j", "s", "Cronometragem"],
            ["h(t)", "Altura da superfície livre ao longo do tempo", "m", "Análise do vídeo"],
            ["T", "Temperatura da água, se disponível", "°C", "Medição direta"],
            ["D", "Diâmetro interno médio do tanque", "m", "Medição direta"],
            ["V", "Volume de água", "m³", "Calibração"],
            ["A_t", "Área efetiva da seção transversal do tanque", "m²", "Geometria ou calibração"],
            ["A(h)", "Área transversal em função da altura", "m²", "Calibração"],
            ["A_o", "Área do orifício", "m²", "Cálculo"],
            ["g", "Aceleração da gravidade", "m/s²", "Valor adotado"],
            ["v_o", "Velocidade de saída no orifício", "m/s", "Modelo teórico"],
            ["Q", "Vazão volumétrica", "m³/s", "Modelo teórico"],
            ["n", "Número de repetições", "-", "Contagem"],
            ["Δt_ideal", "Duração prevista pelo modelo ideal", "s", "Cálculo"],
            ["Δt̄_exp", "Duração experimental média", "s", "Cálculo"],
            ["s_Δt", "Desvio-padrão das durações", "s", "Cálculo"],
            ["C_d", "Coeficiente de descarga experimental", "-", "Cálculo complementar"],
            ["e_r", "Erro relativo percentual", "%", "Cálculo"],
        ],
        [Cm(1.7), Cm(7.0), Cm(2.0), Cm(5.7)],
        subscript_columns={0},
    )
    add_body(
        document,
        "A área do orifício será calculada a partir de seu diâmetro medido. Caso o furo seja "
        "levemente elíptico, serão utilizados dois diâmetros perpendiculares:",
    )
    add_equation(document, r"A_o = \frac{\pi d^2}{4}", [
        math_symbol("A", "o"), " = ", math_fraction(["π", math_superscript("d", "2")], "4")
    ])
    add_equation(document, r"A_o = \frac{\pi d_x d_y}{4}", [
        math_symbol("A", "o"), " = ", math_fraction(["π", math_symbol("d", "x"), math_symbol("d", "y")], "4")
    ])
    add_body(
        document,
        "Para a faixa aproximadamente cilíndrica, a área do tanque poderá ser estimada "
        "geometricamente. Como alternativa mais robusta para uma PET, será utilizada a calibração "
        "pela relação entre variação de volume e variação de altura:",
    )
    add_equation(document, r"A_t = \frac{\pi D^2}{4}", [
        math_symbol("A", "t"), " = ", math_fraction(["π", math_superscript("D", "2")], "4")
    ])
    add_equation(document, r"A_{\mathrm{t,eff}} \approx \frac{\Delta V}{\Delta h}", [
        math_symbol("A", math_normal("t,eff")), " ≈ ", math_fraction(["Δ", "V"], ["Δ", "h"])
    ])
    add_equation(document, r"A(h) = \frac{\mathrm{d}V}{\mathrm{d}h}", [
        math_function("A", "h"), " = ", math_fraction([math_normal("d"), "V"], [math_normal("d"), "h"])
    ])

    add_heading(document, "6. Hipóteses teóricas")
    assumptions = [
        "A água é incompressível.",
        "O escoamento é transiente e será descrito por uma aproximação quase-estacionária: a lei de Torricelli é aplicada a cada instante.",
        "A pressão na superfície livre e na saída é igual à pressão atmosférica.",
        "A velocidade da superfície livre é muito menor que a velocidade no orifício, pois A_t ≫ A_o.",
        "A garrafa permanece vertical e o escoamento ocorre livremente para a atmosfera.",
        "A seção transversal efetiva do tanque é aproximadamente constante na faixa analisada.",
        "A aceleração da gravidade é g = 9,81 m/s².",
        "No modelo ideal inicial, as perdas de carga e a contração do jato são desprezadas.",
        "Em uma análise complementar, os efeitos reais serão representados pelo coeficiente de descarga C_d.",
    ]
    for assumption in assumptions:
        add_bullet(document, assumption)

    add_heading(document, "7. Modelo teórico")
    add_body(
        document,
        "Aplicando a equação de Bernoulli entre a superfície livre e o centro do orifício, "
        "obtém-se a velocidade ideal de saída pela lei de Torricelli:",
    )
    add_equation(document, r"v_o = \sqrt{2gh}", [math_symbol("v", "o"), " = ", math_radical(["2", "g", "h"])])
    add_body(document, "A vazão volumétrica ideal é:")
    add_equation(document, r"Q = A_o\sqrt{2gh}", [
        "Q", " = ", math_symbol("A", "o"), math_radical(["2", "g", "h"])
    ])
    add_body(document, "A conservação de volume no reservatório fornece:")
    add_equation(document, r"-A_t\frac{\mathrm{d}h}{\mathrm{d}t} = A_o\sqrt{2gh}", [
        "−", math_symbol("A", "t"), math_fraction([math_normal("d"), "h"], [math_normal("d"), "t"]),
        " = ", math_symbol("A", "o"), math_radical(["2", "g", "h"])
    ])
    add_body(
        document,
        "Caso a área transversal varie de forma relevante com a altura, deve-se substituir A_t "
        "pela função calibrada A(h) = dV/dh e utilizar a forma integral:",
    )
    add_equation(document, r"\Delta t = \frac{1}{A_o\sqrt{2g}}\int_{h_f}^{h_i}\frac{A(h)}{\sqrt{h}}\,dh", [
        "Δ", "t", " = ",
        math_fraction("1", [math_symbol("A", "o"), math_radical(["2", "g"])]),
        math_nary("∫", math_symbol("h", "f"), math_symbol("h", "i"), [
            math_fraction(math_function("A", "h"), math_radical("h")), " ", math_normal("d"), "h"
        ])
    ])
    add_body(document, "Integrando entre os níveis h_i e h_f, a duração teórica ideal é:")
    add_equation(document, r"\Delta t_{\mathrm{ideal}} = \frac{2A_t}{A_o\sqrt{2g}}\left(\sqrt{h_i}-\sqrt{h_f}\right)", [
        "Δ", math_symbol("t", math_normal("ideal")), " = ",
        math_fraction(["2", math_symbol("A", "t")], [math_symbol("A", "o"), math_radical(["2", "g"])]),
        math_delimiter([math_radical(math_symbol("h", "i")), " − ", math_radical(math_symbol("h", "f"))])
    ])
    add_body(document, "A evolução ideal da altura ao longo do tempo pode ser representada por:")
    add_equation(document, r"h(t) = \left[\sqrt{h_i}-\frac{A_o\sqrt{2g}}{2A_t}t\right]^2", [
        math_function("h", "t"), " = ",
        math_superscript(math_delimiter([
            math_radical(math_symbol("h", "i")), " − ",
            math_fraction([math_symbol("A", "o"), math_radical(["2", "g"])], ["2", math_symbol("A", "t")]),
            "t"
        ], begin="[", end="]"), "2")
    ])
    add_body(
        document,
        "Para considerar a contração do jato e as perdas reais de forma agregada, introduz-se "
        "o coeficiente de descarga C_d:",
    )
    add_equation(document, r"Q = C_dA_o\sqrt{2gh}", [
        "Q", " = ", math_symbol("C", "d"), math_symbol("A", "o"), math_radical(["2", "g", "h"])
    ])
    add_equation(document, r"\Delta t_{\mathrm{mod}}(C_d) = \frac{2A_t}{C_dA_o\sqrt{2g}}\left(\sqrt{h_i}-\sqrt{h_f}\right)", [
        "Δ", math_symbol("t", math_normal("mod")), math_delimiter(math_symbol("C", "d")), " = ",
        math_fraction(
            ["2", math_symbol("A", "t")],
            [math_symbol("C", "d"), math_symbol("A", "o"), math_radical(["2", "g"])]
        ),
        math_delimiter([math_radical(math_symbol("h", "i")), " − ", math_radical(math_symbol("h", "f"))])
    ])
    add_body(document, "Com os tempos medidos, o coeficiente experimental pode ser estimado por:")
    add_equation(document, r"C_d = \frac{\Delta t_{\mathrm{ideal}}}{\overline{\Delta t}_{\mathrm{exp}}}", [
        math_symbol("C", "d"), " = ",
        math_fraction(
            ["Δ", math_symbol("t", math_normal("ideal"))],
            math_subscript(math_bar(["Δ", "t"]), math_normal("exp")),
        )
    ])
    add_body(
        document,
        "A previsão ideal com C_d = 1 será comparada primeiro com os dados. A estimativa de C_d "
        "será apresentada separadamente como caracterização experimental, e não como validação "
        "independente do mesmo conjunto de medições.",
    )
    add_body(
        document,
        "Para o intervalo aproximadamente cilíndrico e com C_d aproximadamente constante, "
        "a representação de √h em função de t deve ser aproximadamente linear:",
    )
    add_equation(document, r"\sqrt{h(t)} = \sqrt{h_i}-\frac{C_dA_o\sqrt{2g}}{2A_t}t", [
        math_radical(math_function("h", "t")), " = ", math_radical(math_symbol("h", "i")), " − ",
        math_fraction(
            [math_symbol("C", "d"), math_symbol("A", "o"), math_radical(["2", "g"])],
            ["2", math_symbol("A", "t")]
        ),
        "t"
    ])

    add_heading(document, "8. Tratamento dos dados")
    add_body(document, "O tempo experimental médio será calculado a partir de n repetições:")
    add_equation(document, r"\overline{\Delta t}_{\mathrm{exp}} = \frac{1}{n}\sum_{j=1}^{n}\Delta t_j", [
        math_subscript(math_bar(["Δ", "t"]), math_normal("exp")), " = ", math_fraction("1", "n"),
        math_nary("∑", ["j", " = ", "1"], "n", ["Δ", math_symbol("t", "j")])
    ])
    add_body(document, "A dispersão das repetições será expressa pelo desvio-padrão amostral:")
    add_equation(document, r"s_{\Delta t} = \sqrt{\frac{1}{n-1}\sum_{j=1}^{n}\left(\Delta t_j-\overline{\Delta t}_{\mathrm{exp}}\right)^2}", [
        math_symbol("s", ["Δ", "t"]), " = ",
        math_radical(
            [
                math_fraction("1", ["n", " − ", "1"]),
                math_nary(
                    "∑",
                    ["j", " = ", "1"],
                    "n",
                    math_superscript(
                        math_delimiter(
                            [
                                "Δ", math_symbol("t", "j"), " − ",
                                math_subscript(math_bar(["Δ", "t"]), math_normal("exp")),
                            ]
                        ),
                        "2",
                    ),
                ),
            ]
        ),
    ])
    add_body(document, "A diferença percentual em relação ao modelo ideal será:")
    add_equation(document, r"e_r = \left|\frac{\overline{\Delta t}_{\mathrm{exp}}-\Delta t_{\mathrm{ideal}}}{\Delta t_{\mathrm{ideal}}}\right|\times100\%", [
        math_symbol("e", "r"), " = ",
        math_delimiter(
            math_fraction(
                [
                    math_subscript(math_bar(["Δ", "t"]), math_normal("exp")),
                    " − ", "Δ", math_symbol("t", math_normal("ideal")),
                ],
                ["Δ", math_symbol("t", math_normal("ideal"))],
            ),
            begin="|",
            end="|",
        ),
        " × ", "100", "%",
    ])
    add_body(document, "Serão preparados:")
    for item in [
        "tabela com as dimensões do aparato;",
        "tabela com pelo menos três repetições e o tempo experimental médio;",
        "comparação entre o tempo teórico ideal e o tempo experimental;",
        "cálculo do erro relativo percentual, do desvio-padrão e estimativa complementar de C_d;",
        "gráfico com as curvas teórica e experimental de h(t);",
        "gráfico de √h em função de t para avaliar a hipótese quase-estacionária;",
        "discussão das diferenças observadas.",
    ]:
        add_bullet(document, item)

    add_heading(document, "9. Fontes de divergência")
    add_body(
        document,
        "Espera-se que o tempo experimental seja superior ao valor teórico ideal. As principais "
        "fontes de divergência a serem discutidas são:",
    )
    for source in [
        "perdas de carga e efeitos viscosos;",
        "contração do jato na saída;",
        "irregularidades nas bordas do orifício;",
        "variação da seção transversal da garrafa;",
        "incerteza na medição do diâmetro do orifício;",
        "tempo de reação no acionamento do cronômetro;",
        "dificuldade de identificar o nível exato da água no vídeo;",
        "limitação da entrada de ar, caso a tampa da garrafa não seja totalmente removida;",
        "inclinação da câmera, paralaxe e oscilações da superfície livre.",
    ]:
        add_bullet(document, source)
    add_body(
        document,
        "A medição cuidadosa do diâmetro do orifício é especialmente importante, pois sua área "
        "é proporcional a d².",
    )

    add_heading(document, "10. Cronograma")
    add_table(
        document,
        ["Período", "Atividade"],
        [
            ["1 de junho de 2026", "Entrega do protocolo experimental"],
            ["2 a 7 de junho de 2026", "Preparação do aparato e testes preliminares"],
            ["8 a 14 de junho de 2026", "Execução dos ensaios e filmagem"],
            ["15 a 18 de junho de 2026", "Extração e processamento dos dados"],
            ["19 a 21 de junho de 2026", "Cálculos, gráficos, relatório e vídeo"],
            ["22 de junho de 2026", "Conclusão da versão final interna"],
            ["23 a 28 de junho de 2026", "Revisão e ajustes finais"],
            ["29 de junho de 2026", "Entrega final oficial"],
        ],
        [Cm(5.1), Cm(11.3)],
    )

    add_heading(document, "11. Resultados esperados")
    add_body(
        document,
        "Espera-se observar uma redução não linear da altura da coluna d'água ao longo do tempo, "
        "em concordância qualitativa com o modelo quase-estacionário de Bernoulli e continuidade. "
        "O gráfico de √h em função do tempo deverá ser aproximadamente linear no intervalo escolhido. "
        "O tempo experimental "
        "deve ser superior ao tempo ideal devido às perdas reais e à contração do jato, efeitos que "
        "serão analisados por meio do erro relativo e do coeficiente de descarga experimental.",
    )

    add_heading(document, "12. Referência preliminar")
    add_body(
        document,
        "MUNSON, B. R. et al. Fundamentos da Mecânica dos Fluidos. Referência indicada para a "
        "fundamentação teórica; a edição utilizada será especificada no relatório científico final.",
    )

    document.save(DOCUMENT_PATH)


if __name__ == "__main__":
    build_protocol()
